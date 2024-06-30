import os
from docx import Document

def check_valid_run(run):
    skipChars = [".", ",", " ", "[", "]"]

    valid_char = False
    for char in run.text:
        if char not in skipChars:
            valid_char = True
    if valid_char:
        return True
    else:
        return False
    

def process_doc(doc_orig):
    # need eliminated by flask upload function. outputFile path is next on the chopping block 
    #doc_path = r'static/files/micropheno_input_text.docx'

    doc = Document(doc_orig)       #later try inputting file rather than saving it to the hard drive.... kind of a no brainer
    newDoc = Document()

    # appending line labels
    last_speaker_B = False
    lines = []
    count = 0


    for para in doc.paragraphs:
        newP = newDoc.add_paragraph('')
        if para.text[0:2] == "((" and para.text[-2:] == "))":
            newP.add_run(para.text)
        else:
            if any((run.bold and check_valid_run(run)) for run in para.runs) and not last_speaker_B:
                newP.add_run(f'B{count + 1}: {para.text}').bold = True
                last_speaker_B = True
                count += 1
            elif any((run.bold and check_valid_run(run)) for run in para.runs) and last_speaker_B:
                newP.add_run(para.text).bold = True
                # count += 1
            elif last_speaker_B:
                newP.add_run(f'A{count + 1}: {para.text}')
                last_speaker_B = False
                count += 1
            else:
                newP.add_run(para.text)
                # count += 1

    return newDoc
    
